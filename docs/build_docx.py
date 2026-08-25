#!/usr/bin/env python3
"""Build README.docx (landscape A4) with the architecture diagram and literature table."""
import re
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x1B, 0x4F, 0x72)
WARN   = RGBColor(0x92, 0x2B, 0x21)
BODY   = "Calibri"

doc = Document()

# ---------- landscape A4 ----------
s = doc.sections[0]
s.orientation = WD_ORIENT.LANDSCAPE
s.page_width, s.page_height = Mm(297), Mm(210)
s.left_margin = s.right_margin = Mm(20)
s.top_margin = s.bottom_margin = Mm(16)

st = doc.styles["Normal"]
st.font.name = BODY
st.font.size = Pt(10)
st.paragraph_format.space_after = Pt(6)


def shade(cell, hexcolor):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def rich(par, text, size=10):
    """Render **bold** and *italic* markers into runs."""
    for tok in re.split(r"(\*\*.+?\*\*|\*[^*]+?\*)", text):
        if not tok:
            continue
        r = par.add_run(tok.strip("*"))
        r.font.size = Pt(size)
        r.font.name = BODY
        r.bold = tok.startswith("**")
        r.italic = tok.startswith("*") and not tok.startswith("**")
    return par


def heading(text, size=15, color=ACCENT, space_before=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = BODY
    return p


def body(text, size=10, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    rich(p, text, size)
    if italic:
        for r in p.runs:
            r.italic = True
    return p


def callout(label, text):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.cell(0, 0)
    shade(c, "FDEDEC")
    p = c.paragraphs[0]
    r = p.add_run(label + " ")
    r.bold = True
    r.font.color.rgb = WARN
    r.font.size = Pt(10)
    rich(p, text, 10)
    doc.add_paragraph()


# ================= TITLE =================
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PhishDriftBench + DASH"); r.bold = True; r.font.size = Pt(24); r.font.color.rgb = ACCENT

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("The Reported–Deployed Gap in Phishing URL Detection:\n"
              "A Time-, Source-, Provenance- and Evasion-Aware Benchmark\n"
              "with Drift-Aware Selective Hardening")
r.font.size = Pt(13)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Final Year Project — Project Definition, Literature Review and Architecture")
r.font.size = Pt(10); r.italic = True

# ================= ABSTRACT =================
heading("1.  Abstract")
body("Published phishing URL detectors routinely report 97–99.5% accuracy, yet these figures are "
     "almost always produced under one evaluation protocol: a random or k-fold split of a single "
     "dataset, balanced at roughly 50% phishing prevalence, with no adversary in the loop. Each of "
     "those choices is individually defensible and collectively misleading. Random splits leak "
     "time, because phishing URLs arrive in campaigns whose near-identical members land in both "
     "train and test partitions. Single-source evaluation hides source-specific artifacts — a "
     "feature can carry *opposite* meaning in different corpora. Balanced test sets conceal the "
     "base-rate problem: a detector with 99% accuracy and a 1% false-positive rate, deployed at a "
     "real prevalence of 1:10⁴, yields precision near 1%.")
body("This project builds **PhishDriftBench**, a four-axis evaluation protocol that measures "
     "detectors across time, across data sources, against rule-based and generative evasion, and "
     "at deployment-realistic class priors. Under this protocol we re-evaluate five representative "
     "architectures drawn from the base papers. We add a **feature-provenance audit** that "
     "separates statically computable lexical features from those requiring a live network lookup, "
     "and quantifies how much reported accuracy is attributable to lookup artifacts rather than "
     "phishing semantics. Finally we propose **DASH** (Drift-Aware Selective Hardening), a "
     "lightweight mitigation combining drift-stability feature weighting, unsupervised drift "
     "detection, a bounded active-labelling budget, and a calibrated abstention band. The project "
     "is **URL-only**: no page is fetched, rendered, hosted or transmitted at any point.")
callout("Status:", "design and protocol complete; experiments not yet run. Every quantitative "
        "claim in the accompanying paper draft is an explicit unfilled placeholder. No results "
        "are reported until they are measured.")

# ================= IMPROVEMENTS =================
heading("2.  Improvements over the base papers")
body("Each improvement names a specific, verified gap in the ten base papers.")

IMPROVEMENTS = [
 ("1. Temporal evaluation instead of random splits",
  "**Gap:** 0 of 10 base papers uses a chronological split. All use random or k-fold partitioning.  "
  "**Improvement:** train on URLs first observed ≤ T, test on disjoint windows at +1, +3, +6 and "
  "+12 months; report a **decay curve**, not a single accuracy. Paper 5 accidentally demonstrates "
  "why — 99.35% on its 2025 data but 89.17% on an older corpus."),
 ("2. True cross-source transfer",
  "**Gap:** only 1 of 10 (X-PHIDE) performs genuine transfer; Paper 5 evaluates on multiple "
  "datasets but **retrains per dataset**, which measures dataset difficulty, not transfer.  "
  "**Improvement:** train once, evaluate on all source pairs plus leave-one-source-out."),
 ("3. Prevalence-corrected reporting",
  "**Gap:** 0 of 10 reports precision at a realistic deployment base rate; all train and test near "
  "50/50.  **Improvement:** report precision, FPR and alerts per 10⁶ URLs at prevalences of 1:10², "
  "1:10³ and 1:10⁴, alongside balanced figures for comparability."),
 ("4. Evasion testing against our own detector",
  "**Gap:** Paper 10 catalogues cloaking and evasion thoroughly; **no method paper evaluates its "
  "own detector against that catalogue.** Paper 4 states outright that homograph attacks were "
  "never systematically evaluated.  **Improvement:** E1 label-preserving transformations from the "
  "cloaking taxonomy (homoglyph/IDN substitution, subdomain and path padding, percent-encoding, "
  "shortener wrapping, TLD swap, hyphenated brand insertion) and E2 URLs from a contemporary "
  "generative model."),
 ("5. Feature-provenance and lookup-dependency audit  (most novel)",
  "**Gap:** no base paper — and no adjacent work we could find — asks *where* a feature's value "
  "comes from or *when* it was resolved.  **Improvement:** classify every feature as **P1** "
  "static-lexical, **P2** lookup-at-inference (WHOIS, DNS, SSL, domain age, shortener expansion) "
  "or **P3** third-party reputation; then measure accuracy with P2/P3 ablated, accuracy when P2 "
  "features are resolved *today* versus at collection time, and recall when the lookup channel is "
  "unavailable or attacker-controlled. The hazard is concrete: Paper 6 reports WHOIS, DNS and "
  "domain-age features over a 549,346-URL corpus containing **only URL strings**. If those lookups "
  "were performed retrospectively they encode *domain mortality* — phishing domains die within "
  "days, legitimate ones persist for years — which any classifier can exploit without learning "
  "anything about phishing. Paper 4, which deliberately avoids third-party features, is the control."),
 ("6. Drift-stability feature scoring",
  "**Gap:** drift is *mentioned* by 4 of 10 papers and *implemented* by none.  **Improvement:** "
  "score each feature by Population Stability Index across time windows and sources plus the "
  "variance of its permutation importance; partition into drift-stable and drift-brittle *before* "
  "deployment. Validation target: the HTTPS inversion X-PHIDE found (100% of legitimate URLs in "
  "PhiUSIIL, 59.2% in their own corpus) should be flagged without access to the second corpus's "
  "labels."),
 ("7. DASH — a mitigation, not just a diagnosis",
  "**Gap:** 0 of 10 proposes any drift mitigation.  **Improvement:** (i) stability-weighted "
  "training; (ii) **unsupervised** drift detection over the prediction-confidence stream, "
  "requiring no labels — the binding constraint in deployment; (iii) a bounded active-labelling "
  "budget (≈2% of the window by uncertainty sampling) on drift alarm; (iv) a calibrated abstention "
  "band that escalates instead of guessing, reported as a coverage–risk curve."),
 ("8. Duplicate-leakage probe (supporting)",
  "Report exact and near-duplicate rates across the train/test boundary for every corpus, and "
  "accuracy before and after deduplication. With k-fold CV over scraped corpora, campaign "
  "duplicates are near-certain."),
]
for title, txt in IMPROVEMENTS:
    heading(title, size=11, color=ACCENT, space_before=8)
    body(txt)

# ================= ARCHITECTURE =================
doc.add_page_break()
heading("3.  System architecture")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture("arch.png", width=Mm(238))
body("The pipeline is strictly URL-only. Layer 1 assembles dated corpora; Layer 2 generates the "
     "four evaluation axes; Layer 3 classifies features by provenance and scores their drift "
     "stability; Layer 4 hosts five baselines re-implemented from the base papers; Layer 5 is the "
     "proposed DASH mitigation; Layer 6 produces deployment-honest reporting. The dashed return "
     "path carries a drift alarm back into DASH for a warm restart.", size=9)

# ================= LITERATURE REVIEW =================
doc.add_page_break()
heading("4.  Literature review")
body("Ten base papers, all published 2023 or later: eight method or empirical studies plus two reviews.")

HEADERS = ["#", "Title", "Authors", "Publication & Year", "Key Findings",
           "Research Gap", "Methodology", "Future Work"]
WIDTHS  = [7, 38, 29, 24, 42, 36, 42, 39]   # mm, sums to 257

ROWS = [
 ["1", "An Effective Detection Approach for Phishing URL Using ResMLP",
  "S. Remya; M. J. Pillai; K. K. Nair; S. R. Subbareddy; Y. Y. Cho",
  "IEEE Access, vol. 12, 2024",
  "98.29% accuracy; residual pipelining outperforms traditional ML baselines on the same features",
  "Blacklists cannot identify dynamic and newly registered URLs; optimal detection accuracy remains an open pursuit",
  "Lexical + sentiment + domain-age features → convolutional and inverted-residual blocks → MLP; Kaggle dataset, 80:20 random split",
  "Address limitations affecting practical viability; explore self-organising networks and online-learning representations"],
 ["2", "On Phishing URL Detection Using Feature Extension",
  "D. He; Z. Liu; X. Lv; S. Chan; M. Guizani",
  "IEEE Internet of Things Journal, vol. 11, no. 24, 2024",
  "Feature extension enriches information-poor URLs; two-layer BERT→CNN improves detection, including on cryptocurrency phishing",
  "URLs are short and carry little signal; blacklist/whitelist methods have inherent limitations",
  "TextRank keyword extraction builds a feature-extension library; BERT embeddings feed a CNN classifier; single real-phishing dataset",
  "Extend to broader transaction-security threats in the blockchain and cryptocurrency ecosystem"],
 ["3", "Evaluating the Impact of Feature Engineering in Phishing URL Detection: A Comparative Study of URL, HTML, and Derived Features",
  "Y. A. Kustiawan; K. I. Ghauth", "IEEE Access, vol. 13, 2025",
  "99.45% (CatBoost); derived features contribute most; systematic comparison across three feature families and ten models",
  "Prior work studies URL or HTML features in isolation; no comparative study of engineered feature sets across models",
  "101,063-URL corpus; URL / HTML / derived feature groups × 10 ML models; single dataset, random split",
  "Cascading architecture — cheap URL analysis first, heavier content and relationship analysis only for borderline cases; craft further derived features"],
 ["4", "PhishOFE: A Novel Machine Learning Framework for Real-Time Phishing URL Detection With Optimized Feature Engineering",
  "Y. A. Kustiawan; K. I. Ghauth", "IEEE Access, vol. 13, 2025",
  "99.48% (CatBoost) with no third-party features, enabling real-time deployment",
  "Prior studies depend on outdated datasets and on third-party services that break in deployment",
  "Optimised feature-engineering pipeline avoiding third-party lookups; CatBoost plus comparison models; single dataset, random split",
  "Homograph-attack detection via character normalisation, visual-similarity metrics and Unicode-aware features — explicitly not evaluated in the paper"],
 ["5", "Enhanced Phishing Detection Approach Using a Layered Model: Domain Squatting and URL Obfuscation Identification and Lexical Feature-Based Classification",
  "R. Goenka; M. Chawla; N. Tiwari", "IEEE Access, vol. 13, 2025",
  "99.35% (XGBoost) at 12.49 ms; but only 89.17% on an older dataset, which the authors attribute to 2025 target brands differing from 2018",
  "Prevailing approaches overlook domain-squatting and URL obfuscation (brand-jacking), which cause most successful phishing",
  "Layer 1: bad-domain features detect brand names in unwanted positions or misspelled forms; Layer 2: lexical-feature ML classifier; own corpus + 6 datasets, retrained per dataset",
  "Extend the brand-jacking feature set beyond conventional lexical features"],
 ["6", "BGL-PhishNet: Phishing Website Detection Using Hybrid Model — BERT, GNN, and LightGBM",
  "S. Remya; M. J. Pillai; B. S. Aparna; S. R. Subbareddy; Y. Y. Cho",
  "IEEE Access, vol. 13, 2025",
  "Hybrid 97.3% accuracy / 97.8% precision / F1 97.3 / ROC-AUC 0.97; individually BERT 90.4%, GNN 88.3%, LightGBM 85.8%",
  "Single models lack flexibility, turnaround time and scalability on large, evolving datasets",
  "BERT (semantic) + GNN (structural) + LightGBM (metadata) with ensemble late fusion; Kaggle corpus of 549,346 URLs; 10-fold cross-validation",
  "Improve GNN memory efficiency; adapt the model to the evolving phishing landscape"],
 ["7", "Staying Ahead of Phishers: A Review of Recent Advances and Emerging Methodologies in Phishing Detection",
  "S. Kavya; D. Sumathi", "Artificial Intelligence Review, vol. 58, no. 50, 2025",
  "Consolidated taxonomy spanning list-based, ML, DL, graph-based and GAN-based detection",
  "No consolidated view of recent deep-learning, graph and generative approaches",
  "Systematic literature review",
  "Dataset diversity, adversarial robustness, interpretability and real-time deployability named as unfilled gaps"],
 ["8", "A Multimodal Phishing Website Detection System Using Explainable Artificial Intelligence Technologies",
  "A. Vulfin; A. Sulavko; V. Vasiliev; A. Minko; A. Kirillova; A. Samotuga",
  "Machine Learning and Knowledge Extraction (MAKE), 2026",
  "F1 0.989 on own data, 0.953 on MTLP; quantifies an 8–15% temporal decay and a 3–9% adversarial drop; SOC deployment scenario on zero-day URLs",
  "Single-modality detectors are brittle and offer no explanation to a human analyst",
  "Late fusion of CatBoost (URL), 1D-CNN (characters), CodeBERT (HTML) and EfficientNet-B7 (screenshot); SHAP and Grad-CAM plus local LLM explanation",
  "Improve explanation quality; broaden zero-day evaluation"],
 ["9", "XGBoost-Based URL Phishing Detection Method With Cross-Dataset Validation (X-PHIDE)",
  "M. Misiek; T. Hyla", "IEEE Access, vol. 14, 2026",
  "97.80% single-dataset → 77.84% cross-dataset average (Custom 90.72, GramBeddings 70.65, PhiUSIIL 72.16); HTTPS inverts meaning across corpora",
  "Dataset shift is “a factor often overlooked in existing literature”; single-dataset results do not transfer",
  "XGBoost over a cross-dataset feature intersection; three distinct datasets; Chrome plug-in for real-time use",
  "Operational challenges of distribution shift; model compression and lightweight deployment"],
 ["10", "Uncovering the Cloak: A Systematic Review of Techniques Used to Conceal Phishing Websites",
  "W. Li; S. Manickam; S. U. A. Laghari; Y.-W. Chong", "IEEE Access, vol. 11, 2023",
  "Taxonomy of server-side and client-side cloaking, 2012–2022; a small number of sophisticated campaigns account for over 89% of attacks",
  "Cloaking and evasion techniques are scattered across the literature with no unified taxonomy",
  "Systematic literature review (SLR)",
  "Strategies targeting sophisticated campaigns; the role of certificate authorities; collaborative data sharing and faster takedown response"],
]

tbl = doc.add_table(rows=1, cols=len(HEADERS))
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.autofit = False

for i, h in enumerate(HEADERS):
    c = tbl.rows[0].cells[i]
    shade(c, "1B4F72")
    c.width = Mm(WIDTHS[i])
    par = c.paragraphs[0]; par.paragraph_format.space_after = Pt(2)
    r = par.add_run(h); r.bold = True; r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.name = BODY

for n, row in enumerate(ROWS):
    cells = tbl.add_row().cells
    for i, val in enumerate(row):
        cells[i].width = Mm(WIDTHS[i])
        if n % 2 == 1:
            shade(cells[i], "F4F6F7")
        par = cells[i].paragraphs[0]
        par.paragraph_format.space_after = Pt(2)
        r = par.add_run(val); r.font.size = Pt(7.5); r.font.name = BODY

doc.add_paragraph()

# ---- summary table ----
heading("What the table shows at a glance", size=11, space_before=10)
SUM = [("Temporal / chronological split", "0 / 10"),
       ("Prevalence-corrected metrics", "0 / 10"),
       ("Genuine cross-source transfer", "1 / 10  (X-PHIDE)"),
       ("Evasion tested against own detector", "0 / 10"),
       ("Drift detection or mitigation implemented", "0 / 10"),
       ("Feature-provenance audit", "0 / 10")]
t2 = doc.add_table(rows=1, cols=2); t2.style = "Table Grid"
for i, h in enumerate(["Evaluation property", "Base papers satisfying it"]):
    c = t2.rows[0].cells[i]; shade(c, "1B4F72"); c.width = Mm(110 if i == 0 else 50)
    r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
for k, v in SUM:
    cs = t2.add_row().cells
    cs[0].width, cs[1].width = Mm(110), Mm(50)
    cs[0].paragraphs[0].add_run(k).font.size = Pt(9)
    rr = cs[1].paragraphs[0].add_run(v); rr.font.size = Pt(9); rr.bold = True
body("Drift is *mentioned* by four papers and *implemented* by none.")

# ================= SCOPE =================
doc.add_page_break()
heading("5.  Scope, feasibility and honesty constraints")
body("**URL-only.** No crawling, rendering, hosting or third-party API calls. This sidesteps the "
     "dependency Papers 4 and 9 complain about, and makes the evasion axis safe: strings are "
     "perturbed or generated inside an offline dataset, never deployed.")
body("**Compute.** Gradient boosting, a small char-CNN and one BERT fine-tune. A laptop suffices "
     "for most of it; requiring no GPU is itself a selling point for a real-time detector.")
body("**Riskiest ingredient: timestamps.** Axis T needs first-observation dates. Where a source "
     "lacks them, dated dataset *releases* serve as a coarse time axis (2023 → 2024 → 2025 → 2026).")
callout("Non-negotiable.", "Every number in every results table must come from an experiment "
        "actually run. The paper draft carries explicit red resultTODO{} placeholders wherever a "
        "measured value belongs. Fabricating results in work submitted for publication is "
        "research misconduct.")
body("**Novelty must be scoped honestly.** Two concurrent works overlap parts of this proposal and "
     "must be cited: **PhreshPhish** (arXiv 2507.10854, 2025 — temporal splits, LSH leakage removal "
     "and base-rate benchmarks, but a preprint, website-based, and silent on feature provenance) "
     "and **Wei et al.** (IEEE Access 2025 — cross-dataset evaluation; found URL length alone "
     "yields 86.13% and 82.54% on two corpora). Claims in this project are stated as findings "
     "about *this corpus of ten papers*, not about the field as a whole.")
body("**Target venue: IEEE Access.** Seven of the ten base papers appear there, it publishes "
     "benchmark-plus-mitigation work of exactly this shape, and it is Scopus/SCIE indexed with a "
     "4–8 week first decision. Fallbacks: IEEE ICC / GLOBECOM / TrustCom for a shorter conference "
     "version; Computers & Security for a slower, no-APC route.")

doc.save("README.docx")
print("wrote README.docx")
